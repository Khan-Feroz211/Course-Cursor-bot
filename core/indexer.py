"""
core/indexer.py
ML document indexer — extraction, chunking, embedding, FAISS index building.
"""
from __future__ import annotations
import os
import hashlib
import logging
import re
from typing import List, Dict, Tuple, Callable, Optional

import numpy as np
import faiss
import pdfplumber
from sentence_transformers import SentenceTransformer
from docx import Document
from openpyxl import load_workbook

from config.config import AppConfig
from security.storage import MetadataStore
from security.integrity import IndexIntegrityChecker
from security.sanitizer import InputSanitizer

logger = logging.getLogger(__name__)

try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logger.info("pytesseract not available — OCR disabled.")


class Indexer:
    """
    Responsible for:
      1. Detecting file changes (MD5 hashing)
      2. Extracting & chunking PDF text
      3. Embedding chunks (SentenceTransformer)
      4. Building / persisting FAISS index
      5. Verifying index integrity on load
    """

    def __init__(self, cfg: AppConfig, store: MetadataStore):
        self.cfg = cfg
        self.store = store
        self.model = SentenceTransformer(cfg.model.name)

    # ── Hashing ───────────────────────────────────────────────────────────────

    def _hash_file(self, path: str) -> str:
        h = hashlib.md5()
        with open(path, "rb") as f:
            for block in iter(lambda: f.read(65536), b""):
                h.update(block)
        return h.hexdigest()

    def _current_hashes(self, folder: str) -> Dict[str, str]:
        base = InputSanitizer.validate_folder(folder)
        return {
            fn: self._hash_file(str(base / fn))
            for fn in os.listdir(base)
            if fn.lower().endswith((".pdf", ".doc", ".docx", ".xlsx", ".xls"))
        }

    def needs_reindex(self, folder: str) -> bool:
        current = self._current_hashes(folder)
        stored = self.store.load_hashes()
        if current != stored:
            logger.info("File changes detected — re-index required.")
            return True
        index_ok = os.path.exists(self.cfg.storage.faiss_index)
        integrity_ok = IndexIntegrityChecker.verify(
            self.cfg.storage.faiss_index, self.cfg.storage.manifest
        )
        if not index_ok or not integrity_ok:
            logger.warning("FAISS index missing or integrity check failed — re-index required.")
            return True
        logger.info("No changes detected — loading existing index.")
        return False

    # ── Extraction ────────────────────────────────────────────────────────────

    def _extract_page_text(self, page) -> str:
        text = page.extract_text() or ""
        if not text.strip() and OCR_AVAILABLE:
            logger.debug("OCR fallback triggered.")
            img = page.to_image().original
            text = pytesseract.image_to_string(img)
        return re.sub(r"\s+", " ", text).strip()

    def _extract_docx_text(self, filepath: str) -> List[Tuple[str, int]]:
        """Extract text from DOCX file. Returns list of (text, page) tuples."""
        try:
            doc = Document(filepath)
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)
            
            # Join all text and return as single page (DOCX doesn't have pages like PDFs)
            full_text = "\n".join(paragraphs)
            return [(full_text, 1)] if full_text.strip() else []
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return []

    def _extract_excel_text(self, filepath: str) -> List[Tuple[str, int]]:
        """Extract text from Excel file. Returns list of (text, sheet_name) tuples."""
        try:
            wb = load_workbook(filepath, data_only=True)
            text_pages = []
            
            for sheet_idx, sheet_name in enumerate(wb.sheetnames, start=1):
                ws = wb[sheet_name]
                rows = []
                for row in ws.iter_rows(values_only=True):
                    # Filter out None values and convert to string
                    row_text = " | ".join(
                        str(cell).strip() for cell in row 
                        if cell is not None
                    )
                    if row_text.strip():
                        rows.append(row_text)
                
                if rows:
                    sheet_text = "\n".join(rows)
                    # Use sheet index as page number for Excel
                    text_pages.append((sheet_text, sheet_idx))
            
            return text_pages if text_pages else []
        except Exception as e:
            logger.error(f"Error extracting Excel text: {e}")
            return []

    def extract_chunks(self, folder: str) -> Tuple[List[str], List[Dict]]:
        base = InputSanitizer.validate_folder(folder)
        chunks, meta = [], []
        
        # Get all document files (PDF, DOC, DOCX, XLSX, XLS)
        doc_files = [f for f in os.listdir(base) if f.lower().endswith((".pdf", ".doc", ".docx", ".xlsx", ".xls"))]

        if not doc_files:
            raise ValueError("No PDF or Word or Excel documents found in the selected folder.")

        for filename in doc_files:
            fp = base / filename
            if not InputSanitizer.is_safe_file(fp, base):
                logger.warning(f"Skipping unsafe file: {filename}")
                continue
            
            file_hash = self._hash_file(str(fp))
            
            try:
                if filename.lower().endswith(".pdf"):
                    # Handle PDF files
                    with pdfplumber.open(str(fp)) as pdf:
                        for page_num, page in enumerate(pdf.pages, start=1):
                            text = self._extract_page_text(page)
                            if not text:
                                continue
                            words = text.split()
                            size = self.cfg.indexing.chunk_size
                            for i in range(0, len(words), size):
                                chunk = " ".join(words[i : i + size])
                                chunks.append(chunk)
                                meta.append({
                                    "file": filename,
                                    "page": page_num,
                                    "chunk_start": i,
                                    "chunk_text": chunk,
                                    "file_hash": file_hash,
                                })
                
                elif filename.lower().endswith((".docx", ".doc")):
                    # Handle DOCX/DOC files
                    text_chunks = self._extract_docx_text(str(fp))
                    for text, page_num in text_chunks:
                        if not text.strip():
                            continue
                        words = text.split()
                        size = self.cfg.indexing.chunk_size
                        for i in range(0, len(words), size):
                            chunk = " ".join(words[i : i + size])
                            chunks.append(chunk)
                            meta.append({
                                "file": filename,
                                "page": page_num,
                                "chunk_start": i,
                                "chunk_text": chunk,
                                "file_hash": file_hash,
                            })
                
                elif filename.lower().endswith((".xlsx", ".xls")):
                    # Handle Excel files
                    text_chunks = self._extract_excel_text(str(fp))
                    for text, sheet_idx in text_chunks:
                        if not text.strip():
                            continue
                        words = text.split()
                        size = self.cfg.indexing.chunk_size
                        for i in range(0, len(words), size):
                            chunk = " ".join(words[i : i + size])
                            chunks.append(chunk)
                            meta.append({
                                "file": filename,
                                "page": sheet_idx,
                                "chunk_start": i,
                                "chunk_text": chunk,
                                "file_hash": file_hash,
                            })
            except Exception as e:
                logger.error(f"Failed to process '{filename}': {e}")

        return chunks, meta

    # ── Index building ────────────────────────────────────────────────────────

    def build(
        self,
        folder: str,
        progress_cb: Optional[Callable[[int, int], None]] = None,
    ) -> faiss.Index:
        chunks, meta = self.extract_chunks(folder)
        if not chunks:
            raise ValueError("No text could be extracted from the documents.")

        total = len(chunks)
        logger.info(f"Embedding {total} chunks…")

        embeddings = self.model.encode(
            chunks,
            batch_size=self.cfg.model.batch_size,
            show_progress_bar=False,
        )

        if progress_cb:
            progress_cb(total // 2, total)

        emb_arr = np.array(embeddings).astype("float32")
        dim = emb_arr.shape[1]

        if total > self.cfg.indexing.ivf_threshold:
            logger.info(f"Large dataset ({total} chunks) — using IVF index.")
            nlist = self.cfg.indexing.nlist
            quantizer = faiss.IndexFlatL2(dim)
            index = faiss.IndexIVFFlat(quantizer, dim, nlist)
            index.train(emb_arr)
        else:
            index = faiss.IndexFlatL2(dim)

        index.add(emb_arr)

        # Persist
        os.makedirs(os.path.dirname(self.cfg.storage.faiss_index), exist_ok=True)
        faiss.write_index(index, self.cfg.storage.faiss_index)
        IndexIntegrityChecker.save_manifest(
            self.cfg.storage.faiss_index, self.cfg.storage.manifest
        )

        # Store metadata safely in SQLite
        self.store.clear_chunks()
        self.store.insert_chunks(meta)
        self.store.save_hashes(self._current_hashes(folder))

        if progress_cb:
            progress_cb(total, total)

        logger.info("Index built and saved successfully.")
        return index

    def load(self) -> faiss.Index:
        if not IndexIntegrityChecker.verify(
            self.cfg.storage.faiss_index, self.cfg.storage.manifest
        ):
            raise IntegrityError("Index integrity check failed. Please rebuild the index.")
        return faiss.read_index(self.cfg.storage.faiss_index)


class IntegrityError(RuntimeError):
    pass
